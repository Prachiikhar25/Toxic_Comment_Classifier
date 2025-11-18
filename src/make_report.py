"""Create a Word (.docx) report for the project using python-docx.

Run with the repository venv python:
  .venv\Scripts\python.exe src\make_report.py

This will write `docs/Toxic_Comment_Classification_Report.docx`.
"""
from docx import Document
from docx.shared import Pt
import os

REPORT_TEXT = {
    'title': 'Toxic Comment Classification System using Deep Learning',
    'author': 'Generated report',
    'abstract': (
        'This report describes a Toxic Comment Classification System designed to detect and label toxic behaviors in user-generated text. '
        'The system uses a two-stage strategy: an exact-match lookup built from curated labeled data to reproduce known labels precisely, '
        'and a supervised classifier to predict toxicity for previously unseen content. While the repository includes a TF-IDF baseline, '
        'this report details a proposed deep-learning (transformer) approach, implementation choices, evaluation results, and deployment considerations.'
    ),
    'introduction': (
        'Online platforms increasingly rely on automated moderation to keep communities safe. Toxic comments — insults, obscenity, threats, identity-based attacks — '
        'degrade user experience and can escalate into real-world harm. Automated classification of toxic comments helps scale moderation and supports human review. '
        'This project implements a demonstrator pipeline and outlines how to extend it with deep-learning techniques for production.'
    ),
    'problem_statement': (
        'Given an input comment, the system must determine membership in one or more classes: toxic, severe_toxic, obscene, threat, insult, identity_hate. '
        'Key difficulties are label imbalance, context sensitivity (sarcasm/irony), lexical variation and obfuscation, and domain shift between data sources.'
    ),
    'objective_goal': (
        'Primary goal: deliver a multi-label classifier with reliable inference and an exact-match fallback that preserves human labels. Secondary goals include providing a user-facing web UI, producing evaluation reports, and designing a path to a deep-learning model for improved performance.'
    ),
    'scope_limitations': (
        'Scope: local training and inference, multi-label classification, web UI demo. Limitations: model generalization is not guaranteed across domains; production scaling, online learning, and reviewer workflows are outside the current scope.'
    ),
    'proposed_solution': (
        'Maintain an exact-match normalized lookup for deterministic label reproduction, and fine-tune a transformer-based model (e.g., BERT/RoBERTa) for probabilistic multi-label classification. Use threshold tuning and class imbalance techniques to improve recall on rare labels.'
    ),
    'implementation_details': (
        'Data format: CSV with columns `id, comment_text, toxic, severe_toxic, obscene, threat, insult, identity_hate`.'
        '\nPreprocessing: lowercase, remove URLs/mentions, normalize punctuation, and produce a cleaned "normalized" field for exact-match lookup. For transformer models, use the corresponding tokenizer (WordPiece/BPE) with truncation/padding.'
        '\nBaseline: TF-IDF vectorizer + One-vs-Rest LogisticRegression (scikit-learn) with a saved model bundle and exact-match lookup. Deep-learning pipeline: Hugging Face Transformers, BCEWithLogitsLoss, AdamW, linear scheduler with warmup, and per-label threshold tuning on validation data.'
    ),
    'results': (
        'Baseline TF-IDF results (example run) show high precision on common labels but low recall on rare categories such as threat and identity_hate. The document includes a sample evaluation table and recommendations to improve performance using deep-learning methods.'
    ),
    'project_structure': (
        'Repository content: requirements.txt, README.md, data/, models/, src/ (preprocess.py, train.py, predict.py, make_report.py), app.py, templates/, static/, docs/. Proposed additions for deep-learning experiments: src/nn_train.py, src/nn_predict.py, notebooks for EDA.'
    ),
    'conclusion': (
        'This project provides a working foundation for toxic comment classification with an emphasis on reproducibility (exact-match) and a path to improved performance via deep learning. Future work should focus on model fine-tuning, explainability, and production readiness.'
    ),
    'appendix': {
        'environment': [
            'Python 3.11',
            'Flask, scikit-learn, pandas, numpy, python-docx',
            'Recommended: PyTorch or TensorFlow and Hugging Face Transformers for deep learning experiments'
        ],
        'training_commands': [
            'python -m src.train',
            'python -m src.nn_train  # (if implemented)'
        ]
    }
}


def add_heading(doc, text, level=1):
    doc.add_heading(text, level=level)


def add_paragraph(doc, text):
    p = doc.add_paragraph(text)
    for run in p.runs:
        run.font.size = Pt(11)


def add_bullets(doc, lines):
    for line in lines:
        p = doc.add_paragraph(line, style='List Bullet')
        for run in p.runs:
            run.font.size = Pt(11)


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Light List Accent 1'
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
    for r in rows:
        row_cells = table.add_row().cells
        for i, v in enumerate(r):
            row_cells[i].text = str(v)
    return table


def make_doc(out_path):
    doc = Document()
    doc.core_properties.title = REPORT_TEXT['title']
    # Cover page
    h = doc.add_heading(REPORT_TEXT['title'], level=0)
    h.alignment = 1
    doc.add_paragraph(f"Author: {REPORT_TEXT.get('author','')}")
    doc.add_paragraph(f"Date: {__import__('datetime').date.today().isoformat()}")
    doc.add_page_break()

    # Table of contents (manual)
    add_heading(doc, 'Table of Contents', level=1)
    toc_lines = [
        'Abstract', 'Introduction', 'Problem statement', 'Objective and Goal',
        'Project scope and Limitation', 'Proposed solution', 'Implementation details',
        'Results', 'Project structure', 'Conclusion', 'Appendix'
    ]
    add_bullets(doc, toc_lines)
    doc.add_page_break()

    add_heading(doc, 'Abstract', level=1)
    add_paragraph(doc, REPORT_TEXT['abstract'])

    # Add a short background subsection to create more content
    add_heading(doc, 'Background and motivation', level=2)
    add_paragraph(doc, 'Online communities thrive when users feel safe. Automated moderation reduces workload for human moderators and provides faster responses to abusive content. The ability to classify and triage toxic comments is therefore central to modern platform safety systems.')

    add_heading(doc, 'Introduction', level=1)
    add_paragraph(doc, REPORT_TEXT['introduction'])

    add_heading(doc, 'Problem statement', level=1)
    add_paragraph(doc, REPORT_TEXT['problem_statement'])

    add_heading(doc, 'Key challenges', level=2)
    add_bullets(doc, [
        'Imbalanced label distribution (rare classes like threat).',
        'Context-dependent toxicity (sarcasm, negation).',
        'Adversarial obfuscation (misspelling, deliberate masking).',
        'Domain shift between training and production data.'
    ])

    add_heading(doc, 'Objective and Goal', level=1)
    add_paragraph(doc, REPORT_TEXT['objective_goal'])

    add_heading(doc, 'Project scope and Limitation', level=1)
    add_paragraph(doc, REPORT_TEXT['scope_limitations'])

    add_heading(doc, 'Assumptions', level=2)
    add_bullets(doc, [
        'Dataset labels are mostly correct and drawn from a similar domain.',
        'Users input single comments (short text) rather than long documents.',
        'Offline retraining is acceptable (no strict online latency constraints).'
    ])

    add_heading(doc, 'Proposed solution', level=1)
    add_paragraph(doc, REPORT_TEXT['proposed_solution'])

    add_heading(doc, 'Implementation details', level=1)
    add_paragraph(doc, REPORT_TEXT['implementation_details'])

    add_heading(doc, 'Preprocessing pipeline (detailed)', level=2)
    add_paragraph(doc, 'The preprocessing pipeline includes the following steps:')
    add_bullets(doc, [
        'Unicode normalization and lowercasing.',
        'URL, email, and @mention removal.',
        'Punctuation normalization and whitespace collapse.',
        'Stopword removal for exact-match normalization (optional).',
        'Tokenization via transformer tokenizer for model input.'
    ])

    add_heading(doc, 'Model training (deep-learning)', level=2)
    add_paragraph(doc, 'Suggested training settings for fine-tuning a transformer:')
    add_table(doc, ['Parameter', 'Suggested value'], [
        ('Model', 'roberta-base or bert-base-uncased'),
        ('Loss', 'BCEWithLogitsLoss'),
        ('Optimizer', 'AdamW'),
        ('Learning rate', '1e-5 to 5e-5'),
        ('Batch size', '8-32 (use gradient accumulation if needed)'),
        ('Epochs', '2-5 (early stopping on val F1)'),
    ])

    doc.add_page_break()

    add_heading(doc, 'Results', level=1)
    add_paragraph(doc, REPORT_TEXT['results'])

    add_heading(doc, 'Baseline evaluation (sample)', level=2)
    # sample table from earlier
    eval_rows = [
        ('toxic', '0.91', '0.59', '0.72', '3056'),
        ('severe_toxic', '0.54', '0.23', '0.32', '321'),
        ('obscene', '0.92', '0.61', '0.74', '1715'),
        ('threat', '0.67', '0.14', '0.22', '74'),
        ('insult', '0.83', '0.49', '0.62', '1614'),
        ('identity_hate', '0.73', '0.15', '0.25', '294'),
    ]
    add_table(doc, ['Label', 'Precision', 'Recall', 'F1-score', 'Support'], eval_rows)

    add_heading(doc, 'Analysis of results', level=2)
    add_paragraph(doc, 'The baseline shows strong precision for many labels but struggles with recall on rarer classes. This indicates the model is conservative and often misses true positives in those categories. Strategies to address this include data augmentation, per-class threshold tuning, and stronger contextual models.')

    doc.add_page_break()

    add_heading(doc, 'Project structure', level=1)
    add_paragraph(doc, REPORT_TEXT['project_structure'])

    add_heading(doc, 'Files and purpose (detailed)', level=2)
    add_bullets(doc, [
        '`src/preprocess.py`: text normalization utilities and exact-match normalization.',
        '`src/train.py`: baseline TF-IDF + logistic regression training and lookup builder.',
        '`src/predict.py`: prediction wrapper that prefers lookup then model.',
        '`app.py`, `templates/`, `static/`: Flask demo UI and assets.',
        '`src/make_report.py`: generates this Word report.'
    ])

    doc.add_page_break()

    add_heading(doc, 'Conclusion', level=1)
    add_paragraph(doc, REPORT_TEXT['conclusion'])

    add_heading(doc, 'Appendix', level=1)
    add_heading(doc, 'Environment & commands', level=2)
    add_paragraph(doc, 'Development environment and quick commands:')
    add_bullets(doc, REPORT_TEXT['appendix']['environment'] + REPORT_TEXT['appendix']['training_commands'])

    add_heading(doc, 'References', level=2)
    add_bullets(doc, [
        'Devlin, J. et al., BERT: Pre-training of Deep Bidirectional Transformers for Language Understanding (2018).',
        'Liu, Y. et al., RoBERTa: A Robustly Optimized BERT Pretraining Approach (2019).',
        'Hugging Face Transformers — https://huggingface.co/transformers/'
    ])

    os.makedirs(os.path.dirname(out_path), exist_ok=True)
    doc.save(out_path)


if __name__ == '__main__':
    out = os.path.abspath(os.path.join(os.path.dirname(__file__), '..', 'docs', 'Toxic_Comment_Classification_Report.docx'))
    make_doc(out)
    print('Wrote report to', out)
